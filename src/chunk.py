import os
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from docx import Document as DocxDocument

def load_docx(file_path):
    doc = DocxDocument(file_path)
    return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

def chunk_documents(folder_path):
    docs = []
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=400,
        chunk_overlap=50
    )

    for file in os.listdir(folder_path):
        if file.endswith(".docx"):
            text = load_docx(os.path.join(folder_path, file))
            chunks = splitter.split_text(text)

            for chunk in chunks:
                docs.append(
                    Document(
                        page_content=chunk,
                        metadata={"source": file}
                    )
                )

    return docs